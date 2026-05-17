"""DOCX export for proposals using python-docx (already a project dependency)."""

from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


FALLBACK_SECTION_LABELS = {
    "executive_summary": "Executive Summary",
    "understanding_requirements": "Understanding of Requirements",
    "proposed_solution": "Proposed Solution / Technical Approach",
    "relevant_experience": "Relevant Experience & Case Studies",
    "team_qualifications": "Team & Qualifications",
    "project_timeline": "Project Timeline",
    "methodology": "Methodology",
    "pricing": "Pricing / Commercial Proposal",
    "why_us": "Why Us",
    "appendix": "Appendix / Supporting Materials",
}

FALLBACK_SECTION_ORDER = [
    "executive_summary", "understanding_requirements", "proposed_solution",
    "relevant_experience", "team_qualifications", "project_timeline",
    "methodology", "pricing", "why_us", "appendix",
]


def _resolve_schema(proposal):
    """Return ordered list of (key, label). Prefer proposal.section_labels (custom),
    then fall back to proposal.sections keys, then to the default 10-section schema."""
    custom = proposal.section_labels or {}
    if custom:
        return list(custom.items())
    sections = proposal.sections or {}
    keys = list(sections.keys()) or FALLBACK_SECTION_ORDER
    return [
        (k, FALLBACK_SECTION_LABELS.get(k, k.replace("_", " ").title()))
        for k in keys
    ]


def generate_docx(proposal) -> BytesIO:
    """Generate a DOCX file from a Proposal instance. Returns a BytesIO buffer."""
    doc = DocxDocument()

    # --- Page setup ---
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # --- Styles ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(8)
    style.paragraph_format.line_spacing = 1.15

    # --- Title page ---
    title = proposal.rfp.title or "Untitled Proposal"
    org_name = proposal.org.name if proposal.org else ""

    # Add some spacing before title
    for _ in range(4):
        doc.add_paragraph()

    # Proposal title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Subtitle line
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Proposal Document")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x99)

    # Org name
    if org_name:
        org_para = doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        org_para.space_before = Pt(24)
        org_run = org_para.add_run(f"Prepared by {org_name}")
        org_run.font.size = Pt(12)
        org_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Status and date
    from datetime import datetime
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.space_before = Pt(12)
    status_label = "Finalized" if proposal.status == "final" else "Draft"
    meta_run = meta_para.add_run(
        f"{status_label}  •  {datetime.now().strftime('%B %d, %Y')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # Page break after cover
    doc.add_page_break()

    # --- Table of Contents header ---
    toc_heading = doc.add_paragraph()
    toc_run = toc_heading.add_run("Table of Contents")
    toc_run.bold = True
    toc_run.font.size = Pt(18)
    toc_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    toc_heading.space_after = Pt(16)

    sections = proposal.sections or {}
    schema = _resolve_schema(proposal)
    for i, (key, label) in enumerate(schema, 1):
        body = (sections.get(key) or "").strip()
        if not body:
            continue
        toc_item = doc.add_paragraph()
        toc_item.paragraph_format.space_after = Pt(4)
        num_run = toc_item.add_run(f"{i}.  ")
        num_run.font.color.rgb = RGBColor(0x66, 0x66, 0x99)
        label_run = toc_item.add_run(label)
        label_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_page_break()

    # --- Proposal sections ---
    for i, (key, label) in enumerate(schema, 1):
        body = (sections.get(key) or "").strip()
        if not body:
            continue

        # Section heading
        heading_para = doc.add_paragraph()
        heading_para.space_before = Pt(18)
        heading_para.space_after = Pt(12)

        num_run = heading_para.add_run(f"{i}.  ")
        num_run.bold = True
        num_run.font.size = Pt(16)
        num_run.font.color.rgb = RGBColor(0x66, 0x66, 0x99)

        label_run = heading_para.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(16)
        label_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # Divider line
        divider = doc.add_paragraph()
        divider.space_after = Pt(8)
        div_run = divider.add_run("─" * 60)
        div_run.font.size = Pt(6)
        div_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xDD)

        # Section body — split by double newlines for paragraphs
        paragraphs = body.split("\n\n")
        for para_text in paragraphs:
            para_text = para_text.strip()
            if not para_text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(para_text)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # --- Footer note ---
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Generated by Draftly  •  {org_name}"
    )
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    footer_run.italic = True

    # Write to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
